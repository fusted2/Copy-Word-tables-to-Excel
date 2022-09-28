import locale

import PySimpleGUI as sg
import docx  # --pip install python-docx--
import openpyxl  # --pip install openpyxl--
from openpyxl import Workbook  # to create blank Workbook
from openpyxl.styles import Font  # to format Excel cells
from openpyxl.styles.borders import Border, Side  # to apply border

sg.theme('LightBlue')

# symbol for collapsible section(s)
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'


# collapsible section function
def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))


# ---------------SPIN > 1 LAYOUT---------------
opened1 = False  # this section is hidden by default
spinSection = [[sg.Radio('Equally split tables to sheets', 'RADIO2', default=True, key='equalSplit')],
               [sg.Text(
                   'Specify number of tables in each sheet. Ex: 6,7,9 tables in sheet #1, #2 and #3, respectively')],
               [sg.Radio('Input:', 'RADIO2', size=(7, 1), default=False), sg.InputText(size=(15, 1), key='splitSheets')]
               ]

# ---------------MAIN LAYOUT---------------
layout = [[sg.Text('© NGUYEN HOANG PHU, 2021', font=(None, 10, 'bold'), text_color='brown')],
          [sg.Text('Source for Excel template:', size=(25, 1), font=(None, 9, 'bold'))],
          [sg.Radio('Existed Excel file:', 'RADIO0', size=(19, 1), default=True), sg.InputText(key='excelTemplate'),
           sg.FileBrowse()],
          [sg.Radio('Create a blank Excel file (only 1 sheet)', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Source for Word document:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
           sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text(SYMBOL_DOWN, enable_events=True, key='-OPEN SEC1-'),
           sg.Text('Copy to 1 sheet or multiple sheets (max = 100):', size=(40, 1), font=(None, 9, 'bold')),
           sg.Spin(values=(list(range(1, 101))), initial_value=1, size=(5, 2), key='sheets')],
          [sg.Text('(*The number of sheets in the Excel template MUST be greater than or equal to this value)',
                   font=(None, 8, 'italic'), text_color='red')],

          # SPIN > 1 LAYOUT
          [collapse(spinSection, False, '-SEC1-')],

          [sg.Text('Optional (take more time):', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=True, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=True, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],
          [sg.Checkbox('Skip tables:', size=(8, 1), key='skip'), sg.InputText(size=(15, 1), key='skipTables')],
          [sg.Checkbox("Print tables' positions (suggest: column Z or #26)", key='position'),
           sg.Spin(values=(list(range(1, 101))), initial_value=26, size=(5, 2), key='positionCol')],

          [sg.Text('Choose folder to save:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(default_text='file.xlsx', key='savedName')],

          [sg.Button('Submit'), sg.Button('Exit')],
          [sg.Output(size=(80, 3))]
          ]


# wordToExcel function
def wordToExcel(dict):
    # load existed Excel file --or-- create a blank Excel file
    # utilize Excel formulas if needed
    if not dict['newWb']:
        wb = openpyxl.load_workbook(dict['excelTemplate'])
    else:
        wb = Workbook()
    # load the Word file
    doc = docx.Document(dict['wordDocument'])

    sheetNames = wb.sheetnames  # create a list of sheets in the Excel file
    sheets = min(int(dict['sheets']), len(sheetNames))  # in case we enter value > existed sheets

    # handle 'splitSheets'
    if not dict['equalSplit']:  # handle input value
        splitSheets = [int(i) for i in dict['splitSheets'].split(',')]
        if len(splitSheets) > sheets == 1:
            splitSheets[0] = sum(splitSheets)
            del (splitSheets[1:])
        else:
            while len(splitSheets) > sheets:
                splitSheets[-2] = splitSheets[-2] + splitSheets[-1]
                del (splitSheets[-1])
        for i in range(len(splitSheets)):  # transform [3,4,6,7] to [2,6,12,19] => from 0 to 19 => 20 tables
            if i == 0:
                splitSheets[i] = splitSheets[i] - 1  # index starts at 0 while min input = 1
            else:
                splitSheets[i] = splitSheets[i] + splitSheets[i - 1]
    else:  # equally split tables
        total = len(doc.tables)
        splitSheets = []
        if sheets > total:
            sheets = total
        if sheets == 1:
            splitSheets += [total]
        else:
            while total / sheets > 0:
                splitSheets += [total // sheets + (total % sheets > 0)]
                if len(splitSheets) == (sheets - 1):
                    splitSheets += [total - sum(splitSheets)]
                    break
        for i in range(len(splitSheets)):  # transform [5,5,5,5] to [4,9,14,19] => from 0 to 19 => 20 tables
            if i == 0:
                splitSheets[i] = splitSheets[i] - 1  # index starts at 0 while min input = 1
            else:
                splitSheets[i] = splitSheets[i] + splitSheets[i - 1]

    # set thousand & decimal separate:
    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8")  # 1.234.567,999
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8")  # 1,234,567.999

    # set border style for Excel cells
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    # (optional) skip some tables
    if dict['skipTables']:
        skipTables = [int(i) for i in dict['skipTables'].split(',')]
        for i in range(len(skipTables)):  # transform [1,3,5,8] (input) to [0,2,4,7] (index)
            skipTables[i] = skipTables[i] - 1
    else:
        skipTables = []

    # reset 'rowNumEx' = 0 when moving to new sheets
    if sheets > 1 and len(splitSheets) > 1:
        reset = []
        r = 0
        while r < len(splitSheets):
            if r == 0:
                a = 0
            else:
                a = splitSheets[r - 1] + 1
            for i in range(a, splitSheets[-1] + 1):
                if i <= splitSheets[r] and i not in skipTables and i not in reset:
                    reset.append(i)
                    r += 1
                    break

    # set the active sheet for pasting
    if sheets == 1:
        rowNumEx = 0
        activeSheet = wb[sheetNames[0]]

    # loop through all tables in the Word file
    for t in range(len(doc.tables)):
        # (optional) skip some tables
        if dict['skip']:
            if t in skipTables:
                continue

        # set active sheets if we split tables
        if sheets > 1 and len(splitSheets) > 1:
            for i in range(len(splitSheets)):
                if t <= splitSheets[i]:
                    activeSheet = wb[sheetNames[i]]
                    break
            if t in reset:
                rowNumEx = 0

        # assign the table need to be copied
        tableDoc = doc.tables[t]

        # get the size of the table
        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        # create a list ._tc to check merged cells in Word
        checkMerge = []

        # loop through all cells in the table
        for i in range(rowNumDoc):
            for j in range(colNumDoc):
                # Excel cell starts at (1, 1)
                excelCell = activeSheet.cell(row=i + 1 + rowNumEx, column=j + 1)

                # Word cell starts at (0, 0) = top left corner of table
                # Convert number from Word to number in Excel
                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  # integer; (4123) = -4123
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  # float; (4.123) = -4.123
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  # string

                # add ._tc to the 'check merged' list
                checkMerge.append(tableDoc.cell(i, j)._tc)

                # (optional) format Excel cells
                if dict['border']:
                    excelCell.border = thin_border  # apply cell border
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  # exclude empty cells in Word
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  # check format in Word: True/False/None
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  # check format in Word: True/False/None
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  # check format in Word: True/False/None
                        except:
                            continue

                        # underline: {'single', 'singleAccounting', 'doubleAccounting', 'double'}
                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))

            # # (optional) just in case we need to know the position of the table copied from Word
            if dict['position']:
                tableNote = activeSheet.cell(row=i + 1 + rowNumEx, column=int(dict['positionCol']))
                tableNote.value = 'This is table %s' % (t + 1)  # t should start at 1 not 0

        # mentioned above: consecutive order and separate by 3 rows
        rowNumEx += rowNumDoc + 3

    # save a new Excel file just in case the code messes everything up
    if not dict['savedName'].endswith('.xlsx'):
        dict['savedName'] += '.xlsx'
    wb.save((dict['savedFolder'] + '/' + dict['savedName']))


# load GUI
window = sg.Window('Copy Word tables to Excel', layout)

while True:  # Event Loop
    event, values = window.read()
    # print(event, values)
    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Submit':
        # Print message to Output box
        print('Running............')
        # convert the 'window dict' to 'normal dict' because 'window dict' values have different class(es)
        keys = ['excelTemplate', 'newWb', 'wordDocument',
                'locale',
                'sheets', 'equalSplit', 'splitSheets',
                'border', 'formatText', 'hyphenToZero', 'skip', 'skipTables', 'position', 'positionCol',
                'savedFolder', 'savedName'
                ]
        convertDict = {x: window[x].get() for x in keys}

        # run 'wordToExcel'
        if convertDict['wordDocument'] and convertDict['savedFolder']:
            wordToExcel(convertDict)
        # Print message to Output box
        print('Done ヅ')

    if event.startswith('-OPEN SEC1-'):
        opened1 = not opened1
        window['-OPEN SEC1-'].update(SYMBOL_UP if opened1 else SYMBOL_DOWN)
        window['-SEC1-'].update(visible=opened1)

window.close()
