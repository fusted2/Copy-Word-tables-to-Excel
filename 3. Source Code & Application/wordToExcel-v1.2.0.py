import datetime
import locale
import logging
import os
import traceback
import PySimpleGUI as sg  # --pip install pysimplegui--
import docx  # --pip install python-docx--
import openpyxl  # --pip install openpyxl--
from openpyxl import Workbook  # to create blank Workbook
from openpyxl.styles import Font  # to format Excel cells
from openpyxl.styles.borders import Border, Side  # to apply border


# THEME & FONT SETUP
# Add your new theme colors and settings
my_new_theme = {'BACKGROUND': '#FEFAF4',
                'TEXT': '#1C3249',
                'INPUT': '#FFFFFF',
                'TEXT_INPUT': '#000000',
                'SCROLL': '#c7e78b',
                'BUTTON': ('black', '#EE8A50'),  # should use 'black' rather than hex code
                'PROGRESS': ('#1C3249', '#EE8A50'),
                'BORDER': 1,
                'SLIDER_DEPTH': 0,
                'PROGRESS_DEPTH': 0
                }
# Add your dictionary to the PySimpleGUI themes
sg.theme_add_new('MyNewTheme', my_new_theme)
# Switch your theme to use the newly added one. You can add spaces to make it more readable
sg.theme('My New Theme')
# Make sure your system have this font first
app_font = 'Open Sauce Sans'


# symbol for collapsible section(s)
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'


# logging debug
logger = logging.getLogger(__name__)


# collapsible section function
def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))


# saved check point function
def saved_check_point(wb, savedFolder, savedName):
    if not savedName.endswith('.xlsx'):
        savedName += '.xlsx'
    wb.save((savedFolder + '/' + savedName))
    return


# ---------------SPIN > 1 LAYOUT---------------
opened1 = True  # this section is visible by default
spinSection = [[sg.Radio('Equally split tables to sheets', 'RADIO2', default=True, key='equalSplit', font=(app_font, 9))],
               [sg.Text(
                   'Specify number of tables in each sheet. Ex: 6,7,9 tables in sheet #1, #2 and #3, respectively',
                   font=(app_font, 9))],
               [sg.Radio('Input (sep ","):', 'RADIO2', size=(15, 1), default=False, font=(app_font, 9)),
                sg.InputText(size=(20, 1), key='splitSheets')]
               ]

# ---------------MAIN LAYOUT---------------
layout = [[sg.Text('© NGUYEN HOANG PHU, 2021', font=(app_font, 10, 'bold'), text_color='brown')],
          [sg.Text('Source for Excel template:', size=(25, 1), font=(app_font, 9, 'bold'))],
          [sg.Radio('Existed Excel file:', 'RADIO0', size=(25, 1), default=True, font=(app_font, 9)),
           sg.InputText(key='excelTemplate'),
           sg.FileBrowse()],
          [sg.Radio('Create a blank Excel file (only 1 sheet)', 'RADIO0', default=False, key='newWb', font=(app_font, 9))],

          [sg.Text('Source for Word document:', size=(25, 1), font=(app_font, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],

          [sg.Text('Thousand & Decimal separator in Word document:', font=(app_font, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale', font=(app_font, 9)),
           sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1), font=(app_font, 9))],

          [sg.Text(SYMBOL_DOWN, enable_events=True, key='-OPEN SEC1-'),
           sg.Text('Copy to 1 sheet or multiple sheets (max = 100):', size=(40, 1), font=(app_font, 9, 'bold')),
           sg.Spin(values=(list(range(1, 101))), initial_value=1, size=(5, 2), key='sheets')],
          [sg.Text('(*The number of sheets in the Excel template MUST be greater than or equal to this value)',
                   font=(app_font, 8, 'italic'), text_color='red')],

          # SPIN > 1 LAYOUT
          [collapse(spinSection, True, '-SEC1-')],

          # OPTIONAL
          [sg.Text('Optional (take more time):', font=(app_font, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=True, key='border', font=(app_font, 9))],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=True, key='formatText', font=(app_font, 9))],
          [sg.Checkbox('Convert cell "-" to 0 (Accounting)', default=False, key='hyphenToZero', font=(app_font, 9))],
          [sg.Checkbox('Skip tables (sep ","):', size=(16, 1), key='skip', font=(app_font, 9)),
           sg.InputText(size=(20, 1), key='skipTables')],
          [sg.Checkbox("Print tables' positions (suggest: column Z or #26)", key='position', font=(app_font, 9)),
           sg.Spin(values=(list(range(1, 101))), initial_value=26, size=(5, 2), key='positionCol')],
          [sg.Checkbox('Saved check point (recommended)', default=True, key='checkPoint', font=(app_font, 9))],

          [sg.Text('Choose location to save output:', size=(25, 1), font=(app_font, 9, 'bold')),
           sg.InputText(key='savedFolder'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(app_font, 9, 'bold')),
           sg.InputText(default_text='file.xlsx', key='savedName')],

          [sg.Button('Submit', font=(app_font, 9))],
          [sg.Output(size=(80, 5))]
          ]


# wordToExcel function
def wordToExcel(dict):
    # load existed Excel file --or-- create a blank Excel file
    # utilize Excel formulas if needed
    if not dict['newWb']:
        wb = openpyxl.load_workbook(dict['excelTemplate'])
    else:
        wb = Workbook()
    # load the Word file
    doc = docx.Document(dict['wordDocument'])

    # sheet check
    sheetNames = wb.sheetnames  # create a list of sheets in the Excel file
    sheets = min(int(dict['sheets']), len(sheetNames))  # in case we enter value > existed sheets

    # handle 'splitSheets'
    if not dict['equalSplit']:  # handle input value
        splitSheets = [int(i) for i in dict['splitSheets'].split(',')]
        if len(splitSheets) > sheets == 1:
            splitSheets[0] = sum(splitSheets)
            del (splitSheets[1:])
        else:
            while len(splitSheets) > sheets:
                splitSheets[-2] = splitSheets[-2] + splitSheets[-1]
                del (splitSheets[-1])
        for i in range(len(splitSheets)):  # transform [3,4,6,7] to [2,6,12,19] => from 0 to 19 => 20 tables
            if i == 0:
                splitSheets[i] = splitSheets[i] - 1  # index starts at 0 while min input = 1
            else:
                splitSheets[i] = splitSheets[i] + splitSheets[i - 1]
    else:  # equally split tables
        total = len(doc.tables)
        splitSheets = []
        if sheets > total:
            sheets = total
        if sheets == 1:
            splitSheets += [total]
        else:
            while total / sheets > 0:
                splitSheets += [total // sheets + (total % sheets > 0)]
                if len(splitSheets) == (sheets - 1):
                    splitSheets += [total - sum(splitSheets)]
                    break
        for i in range(len(splitSheets)):  # transform [5,5,5,5] to [4,9,14,19] => from 0 to 19 => 20 tables
            if i == 0:
                splitSheets[i] = splitSheets[i] - 1  # index starts at 0 while min input = 1
            else:
                splitSheets[i] = splitSheets[i] + splitSheets[i - 1]

    # set thousand & decimal separate:
    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8")  # 1.234.567,999
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8")  # 1,234,567.999

    # set border style for Excel cells
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    # (optional) skip some tables
    if dict['skipTables']:
        skipTables = [int(i) for i in dict['skipTables'].split(',')]
        for i in range(len(skipTables)):  # transform [1,3,5,8] (input) to [0,2,4,7] (index)
            skipTables[i] = skipTables[i] - 1
    else:
        skipTables = []

    # reset 'rowNumEx' = 0 when moving to new sheets
    if sheets > 1 and len(splitSheets) > 1:
        reset = []
        r = 0
        while r < len(splitSheets):
            if r == 0:
                a = 0
            else:
                a = splitSheets[r - 1] + 1
            for i in range(a, splitSheets[-1] + 1):
                if i <= splitSheets[r] and i not in skipTables and i not in reset:
                    reset.append(i)
                    r += 1
                    break

    # set the active sheet for pasting
    if sheets == 1:
        rowNumEx = 0
        activeSheet = wb[sheetNames[0]]

    # loop through all tables in the Word file
    for t in range(len(doc.tables)):
        # (optional) skip some tables
        if dict['skip']:
            if t in skipTables:
                continue

        # set active sheets if we split tables
        if sheets > 1 and len(splitSheets) > 1:
            for i in range(len(splitSheets)):
                if t <= splitSheets[i]:
                    activeSheet = wb[sheetNames[i]]
                    break
            if t in reset:
                rowNumEx = 0

        # assign the table need to be copied
        tableDoc = doc.tables[t]

        # get the size of the table
        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        # create a list ._tc to check merged cells in Word
        checkMerge = []

        # loop through all cells in the table
        for i in range(rowNumDoc):
            for j in range(colNumDoc):
                # Excel cell starts at (1, 1)
                excelCell = activeSheet.cell(row=i + 1 + rowNumEx, column=j + 1)

                # Word cell starts at (0, 0) = top left corner of table
                # Convert number from Word to number in Excel
                try:
                    if tableDoc.cell(i, j)._tc not in checkMerge:
                        try:
                            excelCell.value = locale.atoi(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  # integer; (4123) = -4123
                            logging.debug(
                                'Table: {} | row: {} | col: {} | value: {}'.format(t + 1, i + 1, j + 1,
                                                                                   excelCell.value))
                        except ValueError:
                            try:
                                excelCell.value = locale.atof(
                                    tableDoc.cell(i, j).text.replace('(', '-').replace(')',
                                                                                       ''))  # float; (4.123) = -4.123
                                logging.debug(
                                    'Table: {} | row: {} | col: {} | value: {}'.format(t + 1, i + 1, j + 1,
                                                                                       excelCell.value))
                            except ValueError:
                                if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                    excelCell.value = 0
                                    logging.debug(
                                        'Table: {} | row: {} | col: {} | value: {}'.format(t + 1, i + 1, j + 1,
                                                                                           excelCell.value))
                                else:
                                    excelCell.value = tableDoc.cell(i, j).text  # string
                                    logging.debug(
                                        'Table: {} | row: {} | col: {} | value: {}'.format(t + 1, i + 1, j + 1,
                                                                                           excelCell.value))
                except IndexError:
                    continue

                # add ._tc to the 'check merged' list
                checkMerge.append(tableDoc.cell(i, j)._tc)

                # (optional) format Excel cells
                if dict['border']:
                    excelCell.border = thin_border  # apply cell border
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  # exclude empty cells in Word
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  # check format in Word: True/False/None
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  # check format in Word: True/False/None
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  # check format in Word: True/False/None
                        except:
                            continue

                        # underline: {'single', 'singleAccounting', 'doubleAccounting', 'double'}
                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))

            # # (optional) just in case we need to know the position of the table copied from Word
            if dict['position']:
                tableNote = activeSheet.cell(row=i + 1 + rowNumEx, column=int(dict['positionCol']))
                tableNote.value = 'This is table %s' % (t + 1)  # t should start at 1 not 0

        # mentioned above: consecutive order and separate by 3 rows
        rowNumEx += rowNumDoc + 3

        # save check point after each table
        if dict['checkPoint']:
            saved_check_point(wb, dict['savedFolder'], dict['savedName'])

    # save a new Excel file just in case the code messes everything up
    saved_check_point(wb, dict['savedFolder'], dict['savedName'])


# load GUI
window = sg.Window('Copy Word tables to Excel (1.2.0)', layout)

while True:  # Event Loop
    event, values = window.read()
    # print(event, values)
    if event == sg.WIN_CLOSED:
        break
    if event == 'Submit':
        # Print message to Output box
        print('Running............')
        # convert the 'window dict' to 'normal dict' because 'window dict' values have different class(es)
        keys = ['excelTemplate', 'newWb', 'wordDocument',
                'locale',
                'sheets', 'equalSplit', 'splitSheets',
                'border', 'formatText', 'hyphenToZero', 'skip', 'skipTables', 'position', 'positionCol', 'checkPoint',
                'savedFolder', 'savedName'
                ]
        convertDict = {x: window[x].get() for x in keys}

        # ---BEGIN LOGGING SETUP---
        # log folder
        log_dir = os.path.join(os.path.normpath(os.getcwd() + os.sep + os.pardir), 'wordToExcel_logs')
        is_dir_exist = os.path.exists(log_dir)
        if not is_dir_exist:
            os.makedirs(log_dir)
        # log file
        wordPath, wordName = os.path.split(convertDict['wordDocument'])
        wordName, wordExt = os.path.splitext(wordName)
        dt = str(datetime.datetime.now().strftime("%Y_%m_%d_%H%M%S"))
        log_file = os.path.join(log_dir, 'log_{}_{}.log'.format(dt, wordName))
        # logging config
        logging.basicConfig(level=logging.DEBUG, format='%(asctime)s %(levelname)-8s %(message)s',
                            filename=log_file, filemode='w')
        # delete old log files. Log folder has max 30 log files
        old_log_files = os.listdir(log_dir)
        old_log_files = sorted(old_log_files, reverse=False)  # sort ascending
        if len(old_log_files) >= 30:
            for old_log_file in old_log_files[:2]:
                os.remove(os.path.join(log_dir, old_log_file))
        # ---END LOGGING SETUP---

        # run 'wordToExcel'
        if convertDict['wordDocument'] and convertDict['savedFolder']:
            try:
                wordToExcel(convertDict)
            except Exception as e:  # return to main window after crash
                tb = traceback.format_exc()
                msg = str('Please check again:\n'
                          '  • Excel template (.xlsx) (if ticked)\n'
                          '  • Word document source (.docx)\n'
                          '  • Split input (Ex: 1,2,4) (if ticked)\n'
                          '  • Skip input (Ex: 1,2,4) (if ticked)\n'
                          '  • Output location should not be blank\n'
                          "  • Don't open result file while running\n"
                          '  • Otherwise, it may due to the word file\n'
                          '-----------------------------------------------')
                sg.popup_ok(f'ERROR!', msg, e, background_color='#1C3249', text_color='#E8CC41', button_color='#E89B6F')
                logging.exception('Got exception\n\n')
        
        # Print message to Output box
        print('\nDone ヅ')

    if event.startswith('-OPEN SEC1-'):
        opened1 = not opened1
        window['-OPEN SEC1-'].update(SYMBOL_DOWN if opened1 else SYMBOL_UP)
        window['-SEC1-'].update(visible=opened1)

window.close()
